import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x25, 0x2A)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ClassConnect — Comprehensive Development Summary")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B) # Deep Blue

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Summary of Features, UI Redesigns, Merges & Optimizations (Yesterday to Today)")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x43, 0x38, 0xCA) # Indigo

    doc.add_paragraph() # Spacer

    # Executive Overview
    p_h1 = doc.add_paragraph()
    run_h1 = p_h1.add_run("1. Executive Summary")
    run_h1.font.name = 'Arial'
    run_h1.font.size = Pt(15)
    run_h1.font.bold = True
    run_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    p_body = doc.add_paragraph()
    p_body.paragraph_format.line_spacing = 1.15
    p_body.add_run(
        "Over the past 24-36 hours, the ClassConnect platform underwent major frontend UI redesigns, "
        "new interactive component integrations, backend service merges, routing enhancements, and full multi-tier bug fixes. "
        "All changes were empirically verified with 0 build errors across client (Vite 8 React) and server (Express TypeScript) "
        "and pushed to the origin/yusuf Git repository."
    )

    doc.add_paragraph()

    # Detailed Work Items Table
    p_h2 = doc.add_paragraph()
    run_h2 = p_h2.add_run("2. Detailed Work Completed")
    run_h2.font.name = 'Arial'
    run_h2.font.size = Pt(15)
    run_h2.font.bold = True
    run_h2.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    items = [
        (
            "Full-Screen Hero Redesign",
            "HomePage.jsx",
            "Replaced small container with edge-to-edge high-res Indian student tech workspace image (/assets/hero_students_hq.jpg). Shifted headline text and buttons far left over natural light background space."
        ),
        (
            "Universal Search System",
            "UniversalSearchModal.jsx, FloatingNav.jsx",
            "Built interactive search modal with magnifier navbar button & Ctrl+K shortcut. Added instant fuzzy matching outputting Course Categories Output and Actual Courses Output."
        ),
        (
            "The Ultimate Collections Carousel",
            "CategoryShowcase.jsx",
            "Created 6-category fan-out arch deck card layout with spring-based Framer Motion animations. Hovering lifts cards 38px, scales 1.12x, and straightens with category glow shadow. Added custom AI images."
        ),
        (
            "Compare Your Options Chart",
            "CompareOptionsSection.jsx",
            "Redesigned comparison table with clean borders (border-slate-200) on white canvas. Highlighted ClassConnect in light black (#18181B) with green checkmarks (✔) across 7 simplified criteria."
        ),
        (
            "Batch Zero Results Showcase",
            "StudentBatchResultsShowcase.jsx",
            "Created Batch Zero metric highlight card on left + 2-row horizontal scroll container (2x2 grid) with arrow navigation and 3D hover-to-flip cards on right."
        ),
        (
            "Student Video Testimonials",
            "StudentVideoTestimonialsSection.jsx",
            "Implemented 9:16 mobile portrait video player with play/pause/mute controls, floating stats badges, custom posters, and left/right slider controls."
        ),
        (
            "Merged Testimonials & FAQs",
            "MergedTestimonialsFaqSection.jsx",
            "Unified Student Loved Stories (left column) and Help Center Accordion FAQs (right column) side-by-side with borderless depth edge shadows."
        ),
        (
            "Compact Deep Blue Footer",
            "Footer.jsx",
            "Redesigned footer in signature deep brand blue (#0B132B) with crisp white text, legal compliance links (Privacy, Terms, Refund), and compact spacing."
        ),
        (
            "Git Branch Merge (samir-v2)",
            "Server & Router Modules",
            "Merged 13 commits from samir-v2 into yusuf. Integrated Bunny.net CDN media storage, sequential lecture locking, mandatory Aadhaar KYC, referral link attribution, and certificate verification routes while preserving yusuf UI 100%."
        ),
        (
            "Routing & Filtering Bug Fixes",
            "CourseListPage.jsx, CategoryListPage.jsx",
            "Restored FloatingNav h-20 top spacer. Robustified category matching by slug/name/ID. Fixed invalid /course/ to /courses/ detail route. Added thumbnail image cards to /categories page."
        )
    ]

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ["Feature / Task", "Modified Component(s)", "Key Details & Implementation"]
    col_widths = [Inches(1.8), Inches(1.8), Inches(3.2)]

    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], '1E1B4B')

    # Data Rows
    for row_idx, (feature, component, details) in enumerate(items):
        row_cells = table.add_row().cells
        bg_hex = 'F8F8FC' if row_idx % 2 == 0 else 'FFFFFF'

        for i, val in enumerate([feature, component, details]):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_hex)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            if i == 0:
                run.font.bold = True

    doc.add_paragraph()

    # Git Commits Summary Section
    p_h3 = doc.add_paragraph()
    run_h3 = p_h3.add_run("3. Git Commits & Pushes to origin/yusuf")
    run_h3.font.name = 'Arial'
    run_h3.font.size = Pt(15)
    run_h3.font.bold = True
    run_h3.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    commits = [
        ("Commit 33835ef", "feat(categories): add fan-out arch carousel, category thumbnail images, and fix category routing/filtering"),
        ("Commit 375723f", "chore(cms): remove lingering conflict marker in ManageCmsPage"),
        ("Commit fb7d65b", "merge: integrate samir-v2 backend services, media storage, referrals, and verification while preserving yusuf landing UI"),
        ("Commit 88a113f", "feat(ui): update hero layout, add universal search modal, batch results showcase, video testimonials, option comparison chart, and merged FAQ section")
    ]

    for hash_code, msg in commits:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.15
        r_hash = p.add_run(f"{hash_code}: ")
        r_hash.font.bold = True
        r_hash.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)
        r_msg = p.add_run(msg)

    doc.add_paragraph()

    # Verification Section
    p_h4 = doc.add_paragraph()
    run_h4 = p_h4.add_run("4. Build & System Verification")
    run_h4.font.name = 'Arial'
    run_h4.font.size = Pt(15)
    run_h4.font.bold = True
    run_h4.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    verifications = [
        "Client Build Verification: Vite 8.2.0 production build executed successfully (0 compilation errors, built in 742ms).",
        "Server Build Verification: TypeScript compiler (tsc) executed cleanly with 0 errors.",
        "Git Repository Status: Branch yusuf is fully up to date with origin/yusuf."
    ]

    for v in verifications:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.15
        p.add_run(v)

    # Save document
    file_path = "c:/Users/yusuf/OneDrive/Desktop/Course Selling Website/ClassConnect_Development_Summary.docx"
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == "__main__":
    create_document()
